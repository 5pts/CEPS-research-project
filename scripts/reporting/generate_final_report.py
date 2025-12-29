from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
from pathlib import Path

# Paths
WORKSPACE = Path(r"c:\Users\13926\Desktop\CEPS数据汇总")
CLAUDE_FIG = WORKSPACE / "ClaudeFig"
OUTPUT_FILE = WORKSPACE / "CEPS_Data_Report.docx"

def set_font(run, font_name='SimSun', size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

def create_report():
    doc = Document()
    
    # --- Title Page ---
    title = doc.add_heading('CEPS 数据分析与项目进度汇报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('项目名称：学校社会资本的补偿效应研究\n汇报日期：2025-12-27\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # --- Section 1: Data Variables ---
    doc.add_heading('1. 数据变量说明', level=1)
    doc.add_paragraph('本研究基于 CEPS 第二轮调查数据，核心变量定义如下：')
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['变量名称', '定义/含义', '数据类型', '来源', '计算/备注']
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
    
    vars_data = [
        ('bonding_sc_idx', '纽带型社会资本（师生关系）', '数值型', '学生问卷', '均值合成 (w2c11a, w2c09a, w2c09b)'),
        ('bridging_sc_idx', '桥接型社会资本（同伴氛围）', '数值型', '学生问卷', '均值合成 (w2c12, w2c14a, w2c22b等)'),
        ('ses_idx', '家庭社会经济地位 (SES)', '数值型', '学生/家长问卷', '基于父母教育、职业及家庭资产合成'),
        ('hukou_num', '户籍类型', '分类 (0/1)', '学生问卷 w2a18', '0=城市/非农, 1=农村/农业'),
        ('expect_univ_bin', '升学信心 (教育期望)', '二分变量', '学生问卷 w2c11', '1=期望上大学/普高, 0=其他')
    ]
    
    for v_name, v_def, v_type, v_src, v_calc in vars_data:
        row_cells = table.add_row().cells
        row_cells[0].text = v_name
        row_cells[1].text = v_def
        row_cells[2].text = v_type
        row_cells[3].text = v_src
        row_cells[4].text = v_calc
        
    doc.add_paragraph('\n')

    # --- Section 2: Data Quality ---
    doc.add_heading('2. 数据质量评估', level=1)
    
    doc.add_heading('2.1 完整性 (Completeness)', level=2)
    doc.add_paragraph('原始数据 (N=10750) 中部分核心指标存在较高缺失率（如 w2c11a 缺失约 76%），主要源于问卷跳答设计。经过数据清洗与成对删除（Listwise Deletion），最终纳入分析的有效样本量为 N=1957，满足统计建模的基本要求。')
    
    doc.add_heading('2.2 准确性 (Accuracy)', level=2)
    doc.add_paragraph('通过正态性检验发现，核心连续变量（Bonding SC, Bridging SC, SES）均呈现显著的右偏分布（Skewness > 1），不符合正态分布假设。因此，后续分析建议采用对正态性不敏感的模型（如 Logistic 回归）或非参数方法。')
    
    doc.add_heading('2.3 一致性与时效性', level=2)
    doc.add_paragraph('数据源自 CEPS 第二轮（2014-2015学年），逻辑一致性良好。作为横截面分析，不涉及时间序列逻辑校验。')

    # --- Section 3: Existing Analysis ---
    doc.add_heading('3. 现有数据分析', level=1)
    
    doc.add_heading('3.1 变量分布', level=2)
    doc.add_paragraph('图 1 展示了学校社会资本的分布情况。可见大部分学生的社会资本得分集中在均值附近，但存在右尾拖尾。')
    # Insert Image
    img_path = CLAUDE_FIG / "fig3_social_capital.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(5.0))
        
    doc.add_heading('3.2 教育期望分布', level=2)
    doc.add_paragraph('图 2 展示了样本学生的教育期望分布，整体呈现高期望特征（天花板效应）。')
    img_path = CLAUDE_FIG / "fig1_expectations_dist.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(5.0))

    doc.add_heading('3.3 交互效应探索', level=2)
    doc.add_paragraph('图 3 展示了不同 SES 背景下，师生关系对升学信心的影响。曲线斜率的差异初步验证了“补偿效应”的存在。')
    img_path = CLAUDE_FIG / "Interaction_Bonding_SES.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(5.0))

    # --- Section 4: Current Plan Summary ---
    doc.add_heading('4. 当前计划汇总', level=1)
    
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    headers2 = ['阶段', '工作内容', '产出物', '状态']
    for i, h in enumerate(headers2):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        
    tasks_done = [
        ('数据清洗', '多源数据合并、缺失值处理、变量重构', 'cepsw2studentCN_clean.csv', '已完成'),
        ('描述性统计', '分布分析、缺失率计算、图表绘制', 'figures/*.png', '已完成'),
        ('机制探索', 'ML模型探索非线性与阈值效应', 'Decision_Tree_Thresholds.png', '已完成'),
        ('模型评估', '评估调和模型适用性（结论：不适用）', 'Evaluation_Report.md', '已完成')
    ]
    
    for t_stage, t_content, t_out, t_status in tasks_done:
        row_cells = table2.add_row().cells
        row_cells[0].text = t_stage
        row_cells[1].text = t_content
        row_cells[2].text = t_out
        row_cells[3].text = t_status
        
    doc.add_paragraph('\n')

    # --- Section 5: Subsequent TODO List ---
    doc.add_heading('5. 后续 TODO 清单', level=1)
    
    table3 = doc.add_table(rows=1, cols=5)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    headers3 = ['优先级', '任务名称', '具体内容', '预期交付', '截止时间']
    for i, h in enumerate(headers3):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        
    todos = [
        ('P0', 'HLM 多层模型构建', '实施两层（学生-学校）逻辑回归，检验调节效应', '模型结果表 (Odds Ratios)', 'T+1天'),
        ('P0', '论文主体撰写', '整合背景、方法、结果，撰写完整论文草稿', 'Draft_Paper.docx', 'T+3天'),
        ('P1', '非线性效应补充', '使用 GAM 细化阈值分析，补充到论文讨论部分', 'GAM分析章节', 'T+2天'),
        ('P2', '格式规范化', '参考文献整理、图表排版优化', 'Final_Submission.pdf', 'T+4天')
    ]
    
    for t_prio, t_name, t_desc, t_deliv, t_ddl in todos:
        row_cells = table3.add_row().cells
        row_cells[0].text = t_prio
        row_cells[1].text = t_name
        row_cells[2].text = t_desc
        row_cells[3].text = t_deliv
        row_cells[4].text = t_ddl
        
    doc.add_paragraph('\n风险提示：由于样本量筛选后较小（N=1957），需注意模型过拟合风险，建议使用交叉验证。')

    doc.save(OUTPUT_FILE)
    print(f"Report saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    create_report()
