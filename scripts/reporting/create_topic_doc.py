from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_research_topic_doc(filename="Research_Topic_Definition.docx"):
    doc = Document()
    
    # Title
    title = doc.add_heading('Research Topic 初步定义', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    p = doc.add_paragraph()
    run = p.add_run('“学校能否改变命运？学校社会资本对不同家庭背景学生教育期望的补偿效应研究”\n')
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run('Can Schools Change Destiny? The Compensatory Effect of School Social Capital on Educational Expectations Across Family Backgrounds')
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 1: Bonding SC
    doc.add_heading('形式 A：纽带型资本 (Bonding SC) —— “师生关系”', level=1)
    doc.add_paragraph('定义：学生从教师（特别是班主任）处获得的情感支持、工具性帮助（补课）和信息支持。')
    p = doc.add_paragraph()
    run = p.add_run('中国特色指标：')
    run.bold = True
    p.add_run('老师是否找学生谈心？老师是否在课后提供额外辅导？老师是否严厉（管教型支持）？')

    # Section 2: Bridging SC
    doc.add_heading('形式 B：桥接型资本 (Bridging SC) —— “同伴网络/学校氛围”', level=1)
    doc.add_paragraph('定义：学生所处的同伴环境的质量。')
    p = doc.add_paragraph()
    run = p.add_run('中国特色指标：')
    run.bold = True
    p.add_run('班级里有多少同学计划上大学？周围朋友的学习成绩排名如何？（即“近朱者赤”效应）。')

    # Section 3: Moderator
    doc.add_heading('调节变量：家庭背景 (Family Background)', level=1)
    doc.add_paragraph('“家庭背景”需要区分是经济资本还是制度身份：')
    p = doc.add_paragraph()
    run = p.add_run('具体化方向：')
    run.bold = True
    p.add_run('建议聚焦于“家庭社会经济地位 (SES)” 加上 “户籍状态”。')
    p = doc.add_paragraph()
    run = p.add_run('原因：')
    run.bold = True
    p.add_run('在中国，农村户籍或流动儿童的身份，往往比单纯的父母收入更能决定教育命运。')

    # Section 4: Dependent Variable
    doc.add_heading('3. 因变量：教育期望 (Educational Expectations)', level=1)
    p = doc.add_paragraph()
    run = p.add_run('具体化方向：')
    run.bold = True
    p.add_run('建议界定为 “升学信心”。')
    p = doc.add_paragraph()
    run = p.add_run('核心指标：')
    run.bold = True
    p.add_run('学生是否期望获得 本科学历？或者，学生是否期望考上“普高”（相对于职高）？')
    p = doc.add_paragraph('注：在中国，普职分流是第一个命运分流点，考察“考上普高的期望”比“考博士的期望”更具现实张力。')
    p.style = 'Quote'

    # Section 5: Hypotheses
    doc.add_heading('主假设 H1：学校社会资本与学生升学信心正相关', level=1)
    doc.add_paragraph('├── H1a：师生关系（Bonding SC）越好，升学信心越高')
    doc.add_paragraph('└── H1b：同伴氛围（Bridging SC）越好，升学信心越高')

    doc.add_heading('核心假设 H2：家庭背景调节学校社会资本的效应', level=1)
    doc.add_paragraph('├── H2a：师生关系对低SES学生的效应 > 对高SES学生')
    doc.add_paragraph('├── H2b：同伴氛围对低SES学生的效应 > 对高SES学生')
    doc.add_paragraph('└── H2c：师生关系对农村户籍学生的效应 > 对城市户籍学生')

    doc.add_heading('探索性假设 H3：哪种社会资本的补偿效应更强？', level=1)
    doc.add_paragraph('└── 比较 Bonding SC vs Bridging SC 的补偿效应大小')
    
    doc.add_paragraph('边界锁定：初中阶段（五五分流）')

    # Section 6: Methodology
    doc.add_heading('方法论补充：机器学习探索性分析', level=1)
    doc.add_paragraph('本研究采用机器学习方法作为探索性工具，以检测学校社会资本对教育期望的潜在非线性效应与阈值效应。')
    doc.add_paragraph('1. 首先运用灵活模型（如广义加性模型 GAM 或浅层决策树）可视化并识别非线性特征。')
    doc.add_paragraph('2. 随后通过参数化逼近将所得洞见融入可解释回归模型，确保因果解释仍立足于传统实证框架。')
    doc.add_paragraph('ML主要使用在学校社会资本-教育期望的部分。')

    doc.save(filename)
    print(f"Document saved to {filename}")

if __name__ == "__main__":
    create_research_topic_doc()
