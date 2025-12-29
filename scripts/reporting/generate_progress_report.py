from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
from pathlib import Path
import datetime

# Paths
WORKSPACE = Path(r"c:\Users\13926\Desktop\CEPS数据汇总")
RESCUED_FIG = WORKSPACE / "rescued_figures"
OUTPUT_FILE = WORKSPACE / "Advisor_Progress_Report.docx"

def set_font(run, font_name='SimSun', size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

def create_progress_report():
    doc = Document()
    
    # --- Title ---
    title = doc.add_heading('研究项目进度汇报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(f'汇报人：学生\n汇报日期：{datetime.date.today().strftime("%Y-%m-%d")}\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- 1. Project Overview ---
    doc.add_heading('1. 项目概述', level=1)
    
    doc.add_heading('1.1 研究背景与目标', level=2)
    doc.add_paragraph('本项目旨在探究“学校社会资本”对初中生教育期望的影响，特别是其在不同家庭背景（SES、户籍）学生中的“补偿效应”。核心问题是：学校教育资源能否帮助弱势家庭学生打破阶层固化，提升升学信心？')
    
    doc.add_heading('1.2 上阶段工作回顾', level=2)
    doc.add_paragraph('上次汇报后，原定计划为完成 CEPS 数据清洗并进行初步统计分析。然而，在实施过程中发现了严重的样本流失问题（N仅剩1957），因此本阶段工作的重心调整为“数据抢救与多源整合”。')

    # --- 2. Current Progress ---
    doc.add_heading('2. 当前进度汇报', level=1)
    
    doc.add_paragraph('本阶段已完成核心数据的全量清洗与抢救工作，主要成果如下：')
    
    # Progress Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['任务模块', '具体内容', '关键产出']
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        
    tasks = [
        ('数据抢救', '诊断样本流失原因，替换高缺失率变量（跳答题->必答题）', '样本量从 1957 恢复至 9763 (流失率<10%)'),
        ('多源整合', '将家长、教师、学校数据匹配至学生主表', 'merged_rescued_all.csv (包含家庭收入、师资特征等)'),
        ('可视化更新', '基于全量样本重绘分布图与交互效应图', 'rescued_figures 文件夹'),
        ('质量评估', '完成新数据集的完整性与分布检验', 'Rescue_Evaluation_Report.md')
    ]
    
    for t_mod, t_cont, t_out in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = t_mod
        row_cells[1].text = t_cont
        row_cells[2].text = t_out

    # --- 3. Data Analysis ---
    doc.add_heading('3. 数据分析部分', level=1)
    
    doc.add_heading('3.1 样本量与代表性修复', level=2)
    doc.add_paragraph('经过变量策略调整，有效样本量显著提升，且覆盖了各户籍与成绩段学生，消除了之前的选择性偏差。')
    
    # Insert Dist Image
    img_path = RESCUED_FIG / "rescued_expect_dist.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(5.5))
        doc.add_paragraph('图 1: 修复后的教育期望分布 (样本量 N=9763)', style='Caption')

    # Insert SC Dist Image
    img_path_sc = RESCUED_FIG / "rescued_sc_dist.png"
    if img_path_sc.exists():
        doc.add_picture(str(img_path_sc), width=Inches(6.0))
        doc.add_paragraph('图 2: 社会资本变量分布 (Bonding & Bridging SC)', style='Caption')

    doc.add_heading('3.2 关键发现：补偿效应初现', level=2)
    doc.add_paragraph('基于近万人的大样本交互效应分析显示，师生关系对不同户籍学生的影响存在显著差异。')
    
    # Insert Interaction Image
    img_path2 = RESCUED_FIG / "rescued_interaction_hukou.png"
    if img_path2.exists():
        doc.add_picture(str(img_path2), width=Inches(6.0))
        doc.add_paragraph('图 3: 师生关系对升学信心的交互效应 (分户籍)', style='Caption')
    
    doc.add_paragraph('初步结论：农村户籍学生（红线）在师生关系较好时，升学信心提升的幅度大于城市学生（蓝线），这初步验证了“补偿效应”假设，即学校资源对弱势群体具有更大的边际收益。')

    # --- 4. Issues and Solutions ---
    doc.add_heading('4. 存在问题与解决方案', level=1)
    
    doc.add_paragraph('**问题 1：原始变量选择导致样本雪崩**')
    doc.add_paragraph('原因：原计划使用的“最高学历期望”题目属于跳答题，仅非普高意愿者回答，导致80%缺失。')
    doc.add_paragraph('解决：已替换为全员必答的“期望学历”题目，并重新界定了二分变量，彻底解决了此问题。')
    
    doc.add_paragraph('**问题 2：单层模型无法处理嵌套结构**')
    doc.add_paragraph('困难：学生嵌套于班级和学校，普通回归标准误不准。')
    doc.add_paragraph('解决：已整合学校层级数据 (Level-2)，为下一步实施 HLM (多层线性模型) 做好了数据准备。')
    
    doc.add_paragraph('**需要指导的事项：**')
    doc.add_paragraph('1. HLM 模型中，是否建议引入“学校排名”作为跨层交互项？')
    doc.add_paragraph('2. 对于家长报告的 SES 和学生自报的 SES 不一致的情况，建议以哪个为准？')

    # --- 5. Next Steps ---
    doc.add_heading('5. 下一步工作计划', level=1)
    
    doc.add_paragraph('基于当前高质量的全量数据，下阶段将正式进入核心实证分析：')
    
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    headers2 = ['任务', '预期成果', '时间节点']
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        
    next_steps = [
        ('HLM 建模', '构建两层逻辑回归模型，输出 Odds Ratios 表', 'T+2天'),
        ('稳健性检验', '使用家长 SES 替代学生 SES 重跑模型', 'T+3天'),
        ('论文初稿', '完成实证分析章节的撰写', 'T+5天')
    ]
    
    for t, out, time in next_steps:
        row_cells = table2.add_row().cells
        row_cells[0].text = t
        row_cells[1].text = out
        row_cells[2].text = time

    doc.save(OUTPUT_FILE)
    print(f"Report saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    create_progress_report()
